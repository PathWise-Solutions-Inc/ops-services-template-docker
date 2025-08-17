"""
Document Ingestion Pipeline for LightRAG
Handles various document formats and bulk processing
"""

import os
import logging
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import hashlib
import mimetypes

# Document processing libraries
import pypdf
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup

from .core import get_service
from .config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document formats for ingestion"""
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.html': 'html',
        '.htm': 'html',
        '.json': 'json',
        '.csv': 'csv'
    }
    
    def __init__(self):
        self.service = get_service()
        
    def get_file_hash(self, content: bytes) -> str:
        """Generate SHA256 hash of file content for deduplication"""
        return hashlib.sha256(content).hexdigest()
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"[Page {page_num + 1}]\n{text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
            raise
    
    def extract_text_from_html(self, file_path: Path) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text and preserve some structure
            text = soup.get_text(separator='\n')
            
            # Clean up whitespace
            lines = [line.strip() for line in text.splitlines()]
            lines = [line for line in lines if line]
            
            return "\n".join(lines)
        except Exception as e:
            logger.error(f"Failed to process HTML {file_path}: {e}")
            raise
    
    def extract_text_from_markdown(self, file_path: Path) -> str:
        """Extract and convert markdown to plain text"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            
            return soup.get_text(separator='\n')
        except Exception as e:
            logger.error(f"Failed to process Markdown {file_path}: {e}")
            raise
    
    def process_file(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Process a single file and extract text content
        Returns: (text_content, metadata)
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        file_stats = file_path.stat()
        metadata = {
            "filename": file_path.name,
            "file_path": str(file_path),
            "file_size": file_stats.st_size,
            "modified_time": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            "file_type": self.SUPPORTED_EXTENSIONS[extension],
            "extension": extension
        }
        
        # Extract text based on file type
        if extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        elif extension == '.pdf':
            content = self.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            content = self.extract_text_from_docx(file_path)
        elif extension in ['.html', '.htm']:
            content = self.extract_text_from_html(file_path)
        elif extension == '.md':
            content = self.extract_text_from_markdown(file_path)
        elif extension == '.json':
            with open(file_path, 'r', encoding='utf-8') as f:
                import json
                data = json.load(f)
                content = json.dumps(data, indent=2)
        elif extension == '.csv':
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
                content = "\n".join([",".join(row) for row in rows])
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        # Add content hash for deduplication
        metadata["content_hash"] = self.get_file_hash(content.encode('utf-8'))
        metadata["content_length"] = len(content)
        
        return content, metadata
    
    async def ingest_file(self, file_path: Path) -> Dict[str, Any]:
        """Ingest a single file into LightRAG"""
        try:
            logger.info(f"Processing file: {file_path}")
            
            # Process file to extract content
            content, metadata = self.process_file(file_path)
            
            # Insert into LightRAG
            result = await self.service.insert_document(content, metadata)
            
            if result.get("success"):
                logger.info(f"Successfully ingested: {file_path.name}")
            else:
                logger.error(f"Failed to ingest {file_path.name}: {result.get('error')}")
            
            return {
                "file": str(file_path),
                "success": result.get("success", False),
                "metadata": metadata,
                "error": result.get("error")
            }
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return {
                "file": str(file_path),
                "success": False,
                "error": str(e)
            }
    
    async def ingest_directory(
        self, 
        directory_path: Path,
        recursive: bool = True,
        extensions: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Ingest all supported documents from a directory
        
        Args:
            directory_path: Path to directory
            recursive: Whether to process subdirectories
            extensions: List of extensions to process (None = all supported)
        """
        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"Invalid directory: {directory_path}")
        
        # Get list of files to process
        pattern = "**/*" if recursive else "*"
        files_to_process = []
        
        for file_path in directory_path.glob(pattern):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if extensions:
                    if ext in extensions:
                        files_to_process.append(file_path)
                elif ext in self.SUPPORTED_EXTENSIONS:
                    files_to_process.append(file_path)
        
        logger.info(f"Found {len(files_to_process)} files to process")
        
        # Process files
        results = []
        for i, file_path in enumerate(files_to_process, 1):
            logger.info(f"Processing file {i}/{len(files_to_process)}: {file_path.name}")
            result = await self.ingest_file(file_path)
            results.append(result)
        
        # Summary
        successful = sum(1 for r in results if r["success"])
        failed = len(results) - successful
        
        return {
            "total_files": len(files_to_process),
            "successful": successful,
            "failed": failed,
            "results": results
        }
    
    async def ingest_urls(self, urls: List[str]) -> Dict[str, Any]:
        """
        Ingest content from URLs
        This is a placeholder for web scraping functionality
        """
        results = []
        
        for url in urls:
            # This would need implementation of web scraping
            # For now, return placeholder
            results.append({
                "url": url,
                "success": False,
                "error": "URL ingestion not yet implemented"
            })
        
        return {
            "total_urls": len(urls),
            "successful": 0,
            "failed": len(urls),
            "results": results
        }


class IngestionPipeline:
    """Main ingestion pipeline orchestrator"""
    
    def __init__(self):
        self.processor = DocumentProcessor()
        self.stats = {
            "total_processed": 0,
            "successful": 0,
            "failed": 0,
            "start_time": None,
            "end_time": None
        }
    
    async def run(
        self,
        sources: List[Dict[str, Any]],
        batch_size: int = 10
    ) -> Dict[str, Any]:
        """
        Run the ingestion pipeline
        
        Args:
            sources: List of source configurations
                     [{"type": "file", "path": "..."}, 
                      {"type": "directory", "path": "...", "recursive": true},
                      {"type": "url", "urls": [...]}]
            batch_size: Number of documents to process in parallel
        """
        self.stats["start_time"] = datetime.now()
        all_results = []
        
        for source in sources:
            source_type = source.get("type")
            
            if source_type == "file":
                file_path = Path(source["path"])
                result = await self.processor.ingest_file(file_path)
                all_results.append(result)
                
            elif source_type == "directory":
                dir_path = Path(source["path"])
                recursive = source.get("recursive", True)
                extensions = source.get("extensions")
                result = await self.processor.ingest_directory(
                    dir_path, recursive, extensions
                )
                all_results.extend(result.get("results", []))
                
            elif source_type == "urls":
                urls = source.get("urls", [])
                result = await self.processor.ingest_urls(urls)
                all_results.extend(result.get("results", []))
                
            else:
                logger.warning(f"Unknown source type: {source_type}")
        
        # Update stats
        self.stats["end_time"] = datetime.now()
        self.stats["total_processed"] = len(all_results)
        self.stats["successful"] = sum(1 for r in all_results if r.get("success"))
        self.stats["failed"] = self.stats["total_processed"] - self.stats["successful"]
        
        duration = (self.stats["end_time"] - self.stats["start_time"]).total_seconds()
        
        return {
            "stats": self.stats,
            "duration_seconds": duration,
            "results": all_results
        }


# CLI interface for direct execution
async def main():
    """Main entry point for CLI usage"""
    import argparse
    
    parser = argparse.ArgumentParser(description="LightRAG Document Ingestion Pipeline")
    parser.add_argument("path", help="Path to file or directory")
    parser.add_argument("--recursive", "-r", action="store_true", 
                       help="Process directories recursively")
    parser.add_argument("--extensions", "-e", nargs="+",
                       help="File extensions to process")
    
    args = parser.parse_args()
    
    processor = DocumentProcessor()
    path = Path(args.path)
    
    if path.is_file():
        result = await processor.ingest_file(path)
        print(f"Result: {result}")
    elif path.is_dir():
        result = await processor.ingest_directory(
            path, 
            recursive=args.recursive,
            extensions=args.extensions
        )
        print(f"Processed {result['total_files']} files")
        print(f"Successful: {result['successful']}")
        print(f"Failed: {result['failed']}")
    else:
        print(f"Invalid path: {path}")


if __name__ == "__main__":
    asyncio.run(main())